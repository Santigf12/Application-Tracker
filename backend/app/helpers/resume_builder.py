"""
resume_builder.py
Generates a DOCX resume matching the ODT template style.

Usage from FastAPI:
    from resume_builder import build_resume, convert_to_pdf

    docx_bytes = build_resume(manifest, pool)
    pdf_bytes  = convert_to_pdf(docx_bytes)
"""

import os
import subprocess
import tempfile
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ── Colours from ODT ──────────────────────────────────────────────────────────
COLOUR_NAME = RGBColor(0x00, 0x00, 0x00)      # black
COLOUR_SECTION = RGBColor(0x33, 0x33, 0x33)   # #333333
COLOUR_BODY = RGBColor(0x00, 0x00, 0x00)      # #787878
COLOUR_LINK = RGBColor(0x00, 0x56, 0xA6)      # blue

# ── Page geometry ─────────────────────────────────────────────────────────────
PAGE_W = Inches(8.5)
PAGE_H = Inches(11)
MARGIN_TB = Cm(0.42)
MARGIN_LR = Cm(1.48)

CONTENT_W_TWIPS = int((8.5 - (1.48 / 2.54) * 2) * 1440)


# ─────────────────────────────────────────────────────────────────────────────
# XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_bottom_border(paragraph, color="333333", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_spacing(paragraph, before=0, after=0, line=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before * 20)))
    spacing.set(qn("w:after"), str(int(after * 20)))
    if line is not None:
        spacing.set(qn("w:line"), str(int(line * 20)))
        spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)


def _add_right_tab(paragraph, position_twips: int):
    pPr = paragraph._p.get_or_add_pPr()
    tabs_el = pPr.find(qn("w:tabs"))
    if tabs_el is None:
      tabs_el = OxmlElement("w:tabs")
      pPr.append(tabs_el)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(position_twips))
    tabs_el.append(tab)


def _hyperlink(paragraph, text: str, url: str, colour=COLOUR_LINK):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")

    rpr.append(color_el)
    rpr.append(u)
    r.append(rpr)

    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hl.append(r)
    paragraph._p.append(hl)


# ─────────────────────────────────────────────────────────────────────────────
# Document setup
# ─────────────────────────────────────────────────────────────────────────────

def _setup_document() -> Document:
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = PAGE_W
    sec.page_height = PAGE_H
    sec.top_margin = MARGIN_TB
    sec.bottom_margin = MARGIN_TB
    sec.left_margin = MARGIN_LR
    sec.right_margin = MARGIN_LR

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.font.color.rgb = COLOUR_BODY
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    doc.part.numbering_part
    _add_bullet_numbering(doc)

    return doc


def _add_bullet_numbering(doc):
    numbering_part = doc.part.numbering_part
    numbering = numbering_part._element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), "10")

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")

    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")

    lvl_txt = OxmlElement("w:lvlText")
    lvl_txt.set(qn("w:val"), "•")

    align = OxmlElement("w:lvlJc")
    align.set(qn("w:val"), "left")

    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)

    rPr = OxmlElement("w:rPr")
    rFont = OxmlElement("w:rFonts")
    rFont.set(qn("w:ascii"), "Symbol")
    rFont.set(qn("w:hAnsi"), "Symbol")
    rFont.set(qn("w:hint"), "default")
    rPr.append(rFont)

    lvl.append(start)
    lvl.append(fmt)
    lvl.append(lvl_txt)
    lvl.append(align)
    lvl.append(pPr)
    lvl.append(rPr)
    abstract_num.append(lvl)
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "10")

    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), "10")
    num.append(abstract_ref)

    numbering.append(num)


def _bullet_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "10")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    pPr.append(num_pr)

    _set_spacing(p, before=0, after=0)

    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = COLOUR_BODY


# ─────────────────────────────────────────────────────────────────────────────
# Section building blocks
# ─────────────────────────────────────────────────────────────────────────────

def _add_name_header(doc, personal: dict):
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p_name, before=0, after=2)

    run = p_name.add_run(personal["name"])
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    run.font.bold = False
    run.font.color.rgb = COLOUR_NAME

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p_contact, before=0, after=2)

    parts = [personal["phone"], personal["email"], personal["location"]]
    contact_run = p_contact.add_run(" | ".join(parts) + " | ")
    contact_run.font.size = Pt(11)
    contact_run.font.color.rgb = COLOUR_BODY

    links = personal.get("links", [])
    for i, link in enumerate(links):
        _hyperlink(p_contact, link["label"], link["url"])
        if i < len(links) - 1:
            sep = p_contact.add_run(" | ")
            sep.font.size = Pt(11)
            sep.font.color.rgb = COLOUR_BODY


def _section_heading(doc, title: str):
    p = doc.add_paragraph()
    _set_spacing(p, before=5, after=2)
    _add_bottom_border(p, color="333333", size=6)

    run = p.add_run(title.upper())
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOUR_SECTION


def _add_skills(doc, manifest_groups: list, pool: dict):
    _section_heading(doc, "Skills")

    for group in manifest_groups:
        gdata = pool["skill_groups"][group["id"]]
        skills = [pool["skills"][sid]["name"] for sid in group["skills"]]

        p = doc.add_paragraph()
        _set_spacing(p, before=1, after=1)

        label = p.add_run(f'{gdata["title"]}: ')
        label.font.bold = True
        label.font.size = Pt(11)
        label.font.color.rgb = COLOUR_BODY

        val = p.add_run(", ".join(skills))
        val.font.size = Pt(11)
        val.font.color.rgb = COLOUR_BODY


def _add_education(doc, pool: dict):
    edu_list = pool.get("education", [])
    if not edu_list:
        return

    _section_heading(doc, "Education")

    for edu in edu_list:
        p = doc.add_paragraph()
        _set_spacing(p, before=2, after=0)
        _add_right_tab(p, CONTENT_W_TWIPS)

        deg = p.add_run(edu["degree"])
        deg.font.bold = True
        deg.font.size = Pt(11)
        deg.font.color.rgb = COLOUR_BODY

        date_run = p.add_run(f'\t{edu["date"]}')
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = COLOUR_BODY

        p2 = doc.add_paragraph()
        _set_spacing(p2, before=0, after=2)

        inst = p2.add_run(edu["institution"])
        inst.font.size = Pt(11)
        inst.font.color.rgb = COLOUR_BODY


def _add_experience(doc, manifest_exps: list, pool: dict):
    _section_heading(doc, "Experience")

    for exp in manifest_exps:
        edata = pool["experiences"][exp["id"]]

        p = doc.add_paragraph()
        _set_spacing(p, before=2, after=0)
        _add_right_tab(p, CONTENT_W_TWIPS)

        location = edata.get("location", "") or ""
        title_text = f'{edata["title"]}, {edata["company"]}'
        if location:
            title_text += f'; {location}'

        title_run = p.add_run(title_text)
        title_run.font.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = COLOUR_BODY

        date_run = p.add_run(f'\t{edata["start_date"]} \u2013 {edata["end_date"]}')
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = COLOUR_BODY

        if edata.get("description"):
            p_desc = doc.add_paragraph()
            _set_spacing(p_desc, before=0, after=1)

            desc_run = p_desc.add_run(edata["description"])
            desc_run.font.italic = True
            desc_run.font.size = Pt(11)
            desc_run.font.color.rgb = COLOUR_BODY

        for bid in exp["bullets"]:
            bullet = pool["exp_bullets"][bid]
            _bullet_paragraph(doc, bullet["content"])


def _add_publications(doc, pool: dict):
    publications = pool.get("publications", [])
    if not publications:
        return

    for pub in publications:
        p = doc.add_paragraph()
        _set_spacing(p, before=2, after=0)

        citation = p.add_run(pub["citation"])
        citation.font.size = Pt(11)
        citation.font.color.rgb = COLOUR_BODY

        if pub.get("url"):
            p.add_run(" ")
            _hyperlink(p, pub["url"], pub["url"])


def _add_projects(doc, manifest_projects: list, pool: dict):
    _section_heading(doc, "Technical Projects & Publications")
    _add_publications(doc, pool)

    for proj in manifest_projects:
        if not proj.get("include", True):
            continue

        pdata = pool["projects"][proj["id"]]

        p = doc.add_paragraph()
        _set_spacing(p, before=2, after=0)
        _add_right_tab(p, CONTENT_W_TWIPS)

        name_run = p.add_run(pdata["name"])
        name_run.font.bold = True
        name_run.font.size = Pt(11)
        name_run.font.color.rgb = COLOUR_BODY

        if pdata.get("start_date"):
            date_str = pdata["start_date"]
            if pdata.get("end_date"):
                date_str += f' \u2013 {pdata["end_date"]}'

            date_run = p.add_run(f'\t{date_str}')
            date_run.font.size = Pt(11)
            date_run.font.color.rgb = COLOUR_BODY

        if pdata.get("description"):
            p_desc = doc.add_paragraph()
            _set_spacing(p_desc, before=0, after=1)

            desc_run = p_desc.add_run(pdata["description"])
            desc_run.font.italic = True
            desc_run.font.size = Pt(11)
            desc_run.font.color.rgb = COLOUR_BODY

        for bid in proj["bullets"]:
            bullet = pool["proj_bullets"][bid]
            _bullet_paragraph(doc, bullet["content"])


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def build_resume(manifest: dict, pool: dict) -> bytes:
    """
    Build the resume DOCX from an Ollama manifest + DB pool.

    manifest shape:
        {
            "skill_groups": [{"id": "...", "skills": ["id1", "id2"]}, ...],
            "experiences":  [{"id": "...", "bullets": ["id1", ...]}, ...],
            "projects":     [{"id": "...", "include": True, "bullets": [...]}, ...]
        }

    pool shape:
        {
            "personal": {
                "name": str,
                "phone": str,
                "email": str,
                "location": str,
                "links": [{"label": str, "url": str}]
            },
            "education": [
                {"degree": str, "institution": str, "date": str}
            ],
            "publications": [
                {"citation": str, "url": str}
            ],
            "skill_groups": {id: {"title": str}},
            "skills": {id: {"name": str}},
            "experiences": {
                id: {"title", "company", "location", "start_date", "end_date", "description"}
            },
            "exp_bullets": {id: {"content": str}},
            "projects": {
                id: {"name", "start_date", "end_date", "stack", "links", "description"}
            },
            "proj_bullets": {id: {"content": str}},
        }

    Returns raw DOCX bytes.
    """
    doc = _setup_document()

    _add_name_header(doc, pool["personal"])
    _add_skills(doc, manifest["skill_groups"], pool)
    _add_education(doc, pool)
    _add_experience(doc, manifest["experiences"], pool)
    _add_projects(doc, manifest["projects"], pool)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def convert_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Convert DOCX bytes → PDF bytes via LibreOffice headless.
    Requires LibreOffice in PATH.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "resume.docx")
        pdf_path = os.path.join(tmpdir, "resume.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                docx_path,
                "--outdir",
                tmpdir,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice PDF conversion failed:\n{result.stderr}"
            )

        with open(pdf_path, "rb") as f:
            return f.read()