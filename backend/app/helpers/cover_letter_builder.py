"""
cover_letter_builder.py
Generates a DOCX cover letter using the same visual style as resume_builder.py.
"""

from datetime import date
from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.helpers.resume_builder import (
    _setup_document,
    _add_name_header,
    _set_spacing,
    _add_right_tab,
    _add_bottom_border,
    CONTENT_W_TWIPS,
    COLOUR_BODY,
    COLOUR_SECTION,
)


PERSONAL_INFO = {
    "name": "Santiago Fuentes",
    "phone": "(825) 365-0075",
    "email": "jobs@fuentes.it.com",
    "location": "Calgary, Alberta, T3C 1N6",
    "links": [
        {
            "label": "GitHub",
            "url": "https://github.com/Santigf12",
        },
        {
            "label": "Personal",
            "url": "https://santigf12.github.io/",
        },
    ],
}


def _format_today() -> str:
    """
    Format date as: June 9, 2026
    """
    return date.today().strftime("%B %d, %Y").replace(" 0", " ")


def _add_company_date_row(doc, company: str, date_text: str) -> None:
    """
    Adds a resume-style subtitle/section row:

    Government of Alberta                         June 9, 2026
    ----------------------------------------------------------
    """
    paragraph = doc.add_paragraph()
    _set_spacing(paragraph, before=5, after=4)
    _add_right_tab(paragraph, CONTENT_W_TWIPS)
    _add_bottom_border(paragraph, color="333333", size=6)

    company_run = paragraph.add_run(company)
    company_run.font.name = "Times New Roman"
    company_run.font.bold = True
    company_run.font.size = Pt(12)
    company_run.font.color.rgb = COLOUR_SECTION

    date_run = paragraph.add_run(f"\t{date_text}")
    date_run.font.name = "Times New Roman"
    date_run.font.bold = False
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = COLOUR_BODY


def _add_subject_row(doc, position: str) -> None:
    """
    Adds the optional Re: line.
    """
    paragraph = doc.add_paragraph()
    _set_spacing(paragraph, before=0, after=8)

    label_run = paragraph.add_run("Re: ")
    label_run.font.name = "Times New Roman"
    label_run.font.bold = True
    label_run.font.size = Pt(11)
    label_run.font.color.rgb = COLOUR_BODY

    position_run = paragraph.add_run(position)
    position_run.font.name = "Times New Roman"
    position_run.font.size = Pt(11)
    position_run.font.color.rgb = COLOUR_BODY


def _add_body_paragraph(doc, text: str, before: int = 0, after: int = 8) -> None:
    """
    Adds a normal cover letter paragraph.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(paragraph, before=before, after=after)

    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = COLOUR_BODY


def _clean_paragraphs(content: str) -> list[str]:
    """
    Keeps paragraph breaks and preserves intentional signature line breaks.
    """
    normalized = content.replace("\r\n", "\n").replace("\r", "\n")
    blocks = normalized.split("\n\n")

    paragraphs: list[str] = []

    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if lines:
            paragraphs.append("\n".join(lines))

    return paragraphs


def generate_cover_letter_docx(
    email: str,
    company: str,
    content: str,
    position: str | None = None,
) -> bytes:
    """
    Generates a DOCX cover letter.

    Uses:
    - same resume name/contact header
    - same margins
    - same Times New Roman formatting
    - company left + date right, matching resume subtitle rows

    Returns raw DOCX bytes.
    """
    if not company or not company.strip():
        raise ValueError("Company is required.")

    if not content or not content.strip():
        raise ValueError("Cover letter content is required.")

    personal = {
        **PERSONAL_INFO,
        "email": email.strip() if email and email.strip() else PERSONAL_INFO["email"],
    }

    doc = _setup_document()

    # Same top header as resume.
    _add_name_header(doc, personal)

    # Company on the left, date on the right.
    _add_company_date_row(
        doc=doc,
        company=company.strip(),
        date_text=_format_today(),
    )

    # Optional position row.
    if position and position.strip():
        _add_subject_row(doc, position.strip())

    for paragraph_text in _clean_paragraphs(content):
        _add_body_paragraph(doc, paragraph_text, after=8)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()